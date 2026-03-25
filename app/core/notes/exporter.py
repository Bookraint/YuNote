"""
笔记导出：Markdown / TXT / docx
"""
import re
from pathlib import Path

from app.core.notes.note_manager import NoteManager
from app.core.utils.logger import setup_logger

logger = setup_logger("note_exporter")


class NoteExporter:

    def __init__(self, manager: NoteManager, export_dir: Path):
        self.manager = manager
        self.export_dir = export_dir
        self.export_dir.mkdir(parents=True, exist_ok=True)

    def _safe_filename(self, title: str) -> str:
        return re.sub(r'[\\/:*?"<>|]', "_", title).strip()

    def to_markdown(self, note_id: str) -> Path:
        note = self.manager.get(note_id)
        if not note:
            raise ValueError(f"笔记不存在: {note_id}")
        content = self.manager.get_summary(note_id) or self.manager.get_transcript(note_id)
        fname = self._safe_filename(note.title or note_id) + ".md"
        dest = self.export_dir / fname
        dest.write_text(content, encoding="utf-8")
        logger.info("导出 Markdown: %s", dest)
        return dest

    def to_txt(self, note_id: str) -> Path:
        note = self.manager.get(note_id)
        if not note:
            raise ValueError(f"笔记不存在: {note_id}")
        md = self.manager.get_summary(note_id) or self.manager.get_transcript(note_id)
        plain = _md_to_plain(md)
        fname = self._safe_filename(note.title or note_id) + ".txt"
        dest = self.export_dir / fname
        dest.write_text(plain, encoding="utf-8")
        logger.info("导出 TXT: %s", dest)
        return dest

    def to_docx(self, note_id: str) -> Path:
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            raise ImportError("导出 Word 文件需要安装 python-docx：pip install python-docx")

        note = self.manager.get(note_id)
        if not note:
            raise ValueError(f"笔记不存在: {note_id}")

        md = self.manager.get_summary(note_id) or self.manager.get_transcript(note_id)
        doc = Document()

        for line in md.splitlines():
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- [ ] "):
                doc.add_paragraph(f"☐ {line[6:]}", style="List Bullet")
            elif line.startswith("- [x] ") or line.startswith("- [X] "):
                doc.add_paragraph(f"☑ {line[6:]}", style="List Bullet")
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.strip():
                doc.add_paragraph(line)

        fname = self._safe_filename(note.title or note_id) + ".docx"
        dest = self.export_dir / fname
        doc.save(str(dest))
        logger.info("导出 Word: %s", dest)
        return dest


def _md_to_plain(md: str) -> str:
    """简单去除 Markdown 标记，转为纯文本。"""
    text = re.sub(r"^#{1,6}\s+", "", md, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = re.sub(r"^[-*]\s+", "• ", text, flags=re.MULTILINE)
    text = re.sub(r"^- \[[ xX]\] ", "• ", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
